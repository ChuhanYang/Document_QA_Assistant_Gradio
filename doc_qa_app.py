import fitz  # PyMuPDF
import os
import tempfile
import docx
from pathlib import Path
from PIL import Image
import io
from llama_index.core import SimpleDirectoryReader, VectorStoreIndex, Settings
from llama_index.llms.ollama import Ollama
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
import gradio as gr

# Configure local models
llm = Ollama(model="gemma2:2b", request_timeout=600.0)
embed_model = HuggingFaceEmbedding(model_name="sentence-transformers/all-MiniLM-L6-v2")

Settings.llm = llm
Settings.embed_model = embed_model

index = None  # for global access
chat_history = []  # Q&A session history


# Extract text from PDF
def extract_text_from_pdf(file_path):
    text = ""
    doc = fitz.open(file_path)
    for page in doc:
        text += page.get_text()
    return text


# Extract text from DOCX
def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    return "\n".join([p.text for p in doc.paragraphs])


# Get preview from file
def get_file_preview(file_path, ext):
    if ext == ".pdf":
        doc = fitz.open(file_path)
        page = doc.load_page(0)
        pix = page.get_pixmap()
        image_bytes = pix.tobytes("png")
        try:
            image = Image.open(io.BytesIO(image_bytes))
            image.load()
            return image, None
        except Exception as e:
            print("Image conversion failed:", e)
            return None, "⚠️ Could not preview image."

    elif ext == ".txt":
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                text = "".join([f.readline() for _ in range(20)])
            return None, text
        except Exception as e:
            print("Text preview failed:", e)
            return None, "⚠️ Could not preview text."

    elif ext == ".docx":
        doc = docx.Document(file_path)
        text = "\n".join([p.text for p in doc.paragraphs[:10]])
        return None, text

    return None, None


# Create vector index
def create_index(text):
    with tempfile.NamedTemporaryFile("w+", delete=False, suffix=".txt") as f:
        f.write(text)
        f.flush()
        reader = SimpleDirectoryReader(input_files=[f.name])
        docs = reader.load_data()
    return VectorStoreIndex.from_documents(docs)


# Upload & build index with summary and preview
def upload_and_build(file):
    global index, chat_history
    ext = os.path.splitext(file.name)[1].lower()

    if ext == ".pdf":
        text = extract_text_from_pdf(file.name)
    elif ext == ".txt":
        try:
            text = Path(file.name).read_text(encoding="utf-8", errors="replace")
        except Exception as e:
            return [], f"⚠️ Could not read text file: {e}", None, None, None
    elif ext == ".docx":
        text = extract_text_from_docx(file.name)
    else:
        return [], "Unsupported file format.", None, None, None

    index = create_index(text)
    chat_history.clear()

    summary_prompt = "Summarize this document:\n\n" + text[:2000]
    summary = llm.complete(summary_prompt).text.strip()
    chat_history.append(["📄 Summary", summary])

    preview_img, preview_txt = get_file_preview(file.name, ext)
    return chat_history, "✅ File indexed. Ask anything!", preview_img, preview_txt, None


# Ask question and update chat
def ask_question(question, history):
    global index, chat_history
    if index is None:
        return history + [[question, "⚠️ Please upload a document first."]], None

    query_engine = index.as_query_engine(similarity_top_k=3)
    response = query_engine.query(question).response
    history.append([question, response])
    return history, None


# Save chat history
def download_chat(history):
    temp = tempfile.NamedTemporaryFile(delete=False, suffix=".txt", mode="w", encoding="utf-8")
    for q, a in history:
        temp.write(f"Q: {q}\nA: {a}\n\n")
    temp.close()
    return temp.name


# Gradio UI
with gr.Blocks() as ui:
    gr.Markdown("## 📄 Document Q&A Assistant + File Preview")

    with gr.Row():
        with gr.Column(scale=1):
            file_input = gr.File(label="Upload PDF, TXT, or DOCX", file_types=[".pdf", ".txt", ".docx"])
            upload_btn = gr.Button("📤 Upload and Index")
            status_box = gr.Textbox(label="Status")
            with gr.Tab("Preview"):
                with gr.Accordion("📸 First Page / Text Preview", open=False):
                    preview_image = gr.Image(label="Image Preview", visible=True)
                    preview_text = gr.Textbox(label="Text Preview", lines=10, visible=True)

        with gr.Column(scale=2):
            chatbot = gr.Chatbot(label="Q&A Chat")
            question_input = gr.Textbox(label="Ask a question", placeholder="e.g. What is this document about?")
            ask_btn = gr.Button("Ask")
            download_btn = gr.Button("⬇️ Download Chat")
            download_file = gr.File(visible=True)

    upload_btn.click(upload_and_build, inputs=file_input,
                     outputs=[chatbot, status_box, preview_image, preview_text, download_file])
    ask_btn.click(ask_question, inputs=[question_input, chatbot], outputs=[chatbot, download_file])
    download_btn.click(download_chat, inputs=chatbot, outputs=download_file)

ui.launch()